from fastapi import FastAPI, UploadFile, File
import os
from fastapi.middleware.cors import CORSMiddleware
from io import BytesIO
import re
import pdfplumber
from docx import Document
from sentence_transformers import SentenceTransformer
from rapidfuzz import fuzz
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.preprocessing import normalize
from pymongo import MongoClient
import numpy as np
from bson import ObjectId
from cities import get_canonical_city, is_city_match

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

model = SentenceTransformer("all-MiniLM-L6-v2")

client = MongoClient(os.getenv("MONGO_URL"))
db = client["ITJOBS"]
job_collection = db["jobs"]
skills_collection = db["skills"]


def convert_objectid_to_str(obj):
    """Recursively convert ObjectId to string for JSON serialization"""
    if isinstance(obj, ObjectId):
        return str(obj)
    elif isinstance(obj, dict):
        return {k: convert_objectid_to_str(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [convert_objectid_to_str(item) for item in obj]
    return obj


def load_skill_dict():
    skills = skills_collection.find({}, {"name": 1})
    return [s["name"].lower() for s in skills]

skill_dict = load_skill_dict()


def extract_pdf_text(file_like):
    text = ""
    with pdfplumber.open(file_like) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_docx_text(file_like):
    doc = Document(file_like)
    text = []

    # Đọc paragraph
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)

    # Đọc bảng nếu có
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())

            if row_text:
                text.append(" | ".join(row_text))

    return "\n".join(text)

def extract_text_from_file(file_bytes, filename):
    filename = filename.lower()

    if filename.endswith(".pdf"):
        return extract_pdf_text(BytesIO(file_bytes))

    elif filename.endswith(".docx"):
        return extract_docx_text(BytesIO(file_bytes))

    else:
        raise ValueError("Chỉ hỗ trợ file PDF hoặc DOCX")


def extract_skills(text):
    text_lower = text.lower()
    skills = []
    for skill in skill_dict:
        if skill in text_lower or fuzz.partial_ratio(skill, text_lower) > 85:
            skills.append(skill)
    return list(set(skills))


def normalize_text(text):
    if not text:
        return ""
    return re.sub(r"\s+", " ", text.lower().strip())


def extract_city(text):
    """Extract city name from text using cities.py"""
    if not text:
        return ""
    return get_canonical_city(text)


def extract_location_text(text):
    """
    Extract full location/address from CV text.
    """
    if not text:
        return ""

    lines = [line.strip() for line in text.splitlines() if line.strip()]

    # 1. Priority: keyword-based
    markers = [
        "nơi ở",
        "noi o",
        "địa chỉ",
        "dia chi",
        "đ/c",
        "dc",
        "address",
        "location",
        "📍",
        "pin",
    ]

    for line in lines:
        lower = line.lower()

        for marker in markers:
            if marker in lower:

                if marker == "📍":
                    result = line.split(marker, 1)[1].strip()
                else:
                    parts = re.split(re.escape(marker), line, flags=re.IGNORECASE)
                    result = parts[1].strip() if len(parts) > 1 else line.strip()

                result = re.sub(r"^[\-:\s]+", "", result)

                if len(result) > 5:
                    return result

    # 2. Detect real address patterns
    address_patterns = [
        # US style:
        # 50 GRAHAM ST, JERSEY CITY, NJ 07307
        r"\d+\s+[A-Za-z0-9\s\.\-]+,\s*[A-Za-z\s]+,\s*[A-Z]{2}\s+\d{4,6}",

        # Vietnam style:
        # 123 Nguyễn Văn Cừ, Quận 5, TP.HCM
        r"\d+\s+[\w\sÀ-ỹ\.\-]+,\s*.*(?:quận|q\.|huyện|tp\.?|thành phố|district|city)",

        # Generic address line with number + street
        r"\d+\s+.*(?:street|st\.?|road|rd\.?|avenue|ave\.?|boulevard|blvd\.?|lane|ln\.?)",
    ]

    for line in lines:

        # Skip emails/phones
        if "@" in line:
            continue

        digit_count = sum(c.isdigit() for c in line)

        # Address thường phải có số
        if digit_count == 0:
            continue

        for pattern in address_patterns:
            if re.search(pattern, line, re.IGNORECASE):

                # tránh match quá dài
                if len(line) < 120:
                    return line.strip()

    # 3. Fallback by city detection
    for line in lines:
        city = extract_city(line)

        if city:
            return line.strip()

    return ""


def parse_experience_text(text):
    if not text:
        return None

    normalized = text.lower()
    no_experience_markers = [
        "no formal work experience",
        "no experience",
        "chưa có kinh nghiệm",
        "không có kinh nghiệm",
        "no working experience",
        "no professional experience",
    ]
    for marker in no_experience_markers:
        if marker in normalized:
            return 0

    def parse_experience_line(line):
        if not line:
            return None
        line_lower = line.lower()
        range_match = re.search(
            r"(\d+)\s*(?:\+|\-\s*|to|đến|den|–)\s*(\d+)\s*(?:năm|nam|year|years|yr|yrs)",
            line_lower,
        )
        if range_match:
            return max(int(range_match.group(1)), int(range_match.group(2)))
        single_match = re.search(r"(\d+)\s*(?:\+)?\s*(?:năm|nam|year|years|yr|yrs)", line_lower)
        if single_match:
            return int(single_match.group(1))
        return None

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    experience_lines = []
    for line in lines:
        lower_line = line.lower()
        if any(marker in lower_line for marker in [
            "kinh nghiệm",
            "experience",
            "work experience",
            "years of experience",
            "years experience",
            "năm kinh nghiệm",
        ]):
            experience_lines.append(line)

    if not experience_lines:
        for index, line in enumerate(lines):
            lower_line = line.lower()
            if any(marker in lower_line for marker in ["kinh nghiệm", "experience"]):
                if index + 1 < len(lines):
                    experience_lines.append(lines[index + 1])

    for line in experience_lines:
        for marker in no_experience_markers:
            if marker in line.lower():
                return 0
        parsed = parse_experience_line(line)
        if parsed is not None:
            return parsed

    return None


def parse_job_experience(experience):
    if experience is None:
        return None
    if isinstance(experience, (int, float)):
        return int(experience)
    text = str(experience).lower()
    if not text:
        return None
    range_match = re.search(r"(\d+)\s*(?:\+|\-\s*|to|đến|den|–)\s*(\d+)\s*(?:năm|nam|year|years|yr|yrs)?", text)
    if range_match:
        return max(int(range_match.group(1)), int(range_match.group(2)))
    single_match = re.search(r"(\d+)\s*(?:\+)?\s*(?:năm|nam|year|years|yr|yrs)", text)
    if single_match:
        return int(single_match.group(1))
    digits = re.findall(r"\d+", text)
    return int(digits[0]) if digits else None


def get_job_location_city(job):
    """Extract city from job location fields"""
    parts = []
    location_field = job.get("location")
    if isinstance(location_field, dict):
        parts.append(location_field.get("name", ""))
    elif isinstance(location_field, str):
        parts.append(location_field)
    parts.append(job.get("work_location_detail", ""))
    combined = " ".join([p for p in parts if p])
    return extract_city(combined)


def check_location_match(cv_location, job_location):
    """Check if CV location and job location are in same city"""
    if not cv_location or not job_location:
        return False
    # Use cities.py logic for city matching
    return is_city_match(cv_location, job_location)


def generate_job_embedding(text):
    """Generate embedding from text when job doesn't have pre-computed embedding"""
    try:
        emb_full = model.encode(text)
        return emb_full.tolist()
    except:
        return np.zeros(384).tolist()


def prepare_job_skills(job):
    """Extract job skills with fallback for old jobs without structured fields"""
    # Prioritize structured fields
    must_have = job.get("mustHaveSkills", [])
    optional = job.get("optionalSkills", [])
    domain = job.get("domainKnowledge", [])
    
    # Fallback: extract from requirements text if structured fields are empty
    if not must_have and not optional and not domain:
        requirements_text = " ".join(job.get("requirements", []))
        if not requirements_text:
            requirements_text = job.get("description", "")
        
        if requirements_text:
            extracted = extract_skills(requirements_text)
            # Treat all extracted skills as must-have for old jobs
            must_have = extracted
    
    return {
        "must_have": set([s.lower() for s in must_have]) if must_have else set(),
        "optional": set([s.lower() for s in optional]) if optional else set(),
        "domain": set([s.lower() for s in domain]) if domain else set()
    }


@app.post("/recommend")
async def recommend_jobs(file: UploadFile = File(...)):
    try:
        allowed_extensions = [".pdf", ".docx"]

        if not any(file.filename.lower().endswith(ext) for ext in allowed_extensions):
            return {
                "error": "File không hợp lệ. Vui lòng tải lên file CV định dạng PDF hoặc DOCX.",
                "skills_found": [],
                "recommendations": []
            }
        
        # Extract text
        file_bytes = await file.read()
        text = extract_text_from_file(file_bytes=file_bytes, filename=file.filename)
        if not text or text.strip() == "":
            return {
                "error": "Không thể đọc nội dung file. Vui lòng kiểm tra file CV.",
                "skills_found": [],
                "recommendations": []
            }
        
        cv_skills = extract_skills(text)
        cv_skills_set = set(cv_skills)

        # Embedding cho CV
        if len(cv_skills) == 0:
            emb_skill = np.zeros(384)  # all-MiniLM-L6-v2 has 384 dimensions
        else:
            emb_skill = model.encode(" ".join(cv_skills))
        
        emb_full = model.encode(text)
        cv_emb = 0.7 * emb_skill + 0.3 * emb_full

        # Lấy tất cả embedding job kèm requirements/skills
        jobs = list(job_collection.find({}, {
            "title": 1, 
            "embedding": 1,
            "mustHaveSkills": 1,
            "optionalSkills": 1,
            "domainKnowledge": 1,
            "location": 1,
            "salary_raw": 1,
            "company": 1,
            "requirements": 1,
            "description": 1,
            "experience": 1,
            "experienceLevel": 1,
            "work_location_detail": 1
        }))

        if len(jobs) == 0:
            return {
                "skills_found": cv_skills,
                "recommendations": [],
                "message": "Không có công việc nào trong hệ thống"
            }

        # Build job vectors with fallback for missing embeddings
        job_vectors = []
        job_embeddings_fallback = []  # Track which jobs used fallback embeddings
        
        for i, job in enumerate(jobs):
            embedding = job.get("embedding")
            if embedding:
                job_vectors.append(embedding)
                job_embeddings_fallback.append(False)
            else:
                # Fallback: generate embedding from text
                requirements_text = " ".join(job.get("requirements", []))
                if not requirements_text:
                    requirements_text = job.get("description", "")
                
                if requirements_text:
                    generated_emb = generate_job_embedding(requirements_text)
                    job_vectors.append(generated_emb)
                else:
                    job_vectors.append(np.zeros(384).tolist())
                job_embeddings_fallback.append(True)
        
        job_vectors = np.array(job_vectors)

        cv_emb = normalize([cv_emb])[0]
        job_vectors = normalize(job_vectors)
        
        # Tính similarity
        scores = cosine_similarity([cv_emb], job_vectors)[0]

        # Format kết quả với chi tiết matched skills
        cv_full_location = extract_location_text(text)
        cv_city = extract_city(cv_full_location) if cv_full_location else ""
        cv_experience_years = parse_experience_text(text)

        results = []
        for job, score, used_fallback_emb in zip(jobs, scores, job_embeddings_fallback):
            # Prepare job skills with fallback logic
            job_skills = prepare_job_skills(job)
            job_required_skills = job_skills["must_have"]
            job_optional_skills = job_skills["optional"]
            job_domain = job_skills["domain"]
            job_all_skills = job_required_skills | job_optional_skills | job_domain
            
            matched_required = cv_skills_set & job_required_skills
            matched_optional = cv_skills_set & job_optional_skills
            matched_domain = cv_skills_set & job_domain
            matched_total = cv_skills_set & job_all_skills
            
            # Tính tỷ lệ match skill
            skill_score = 0
            if len(job_all_skills) > 0:
                skill_score = (len(matched_total) / len(job_all_skills)) * 100

            # Kinh nghiệm
            job_experience_years = parse_job_experience(job.get("experience") or job.get("experienceLevel"))
            experience_match = None
            experience_score = 50  # default: chưa xác định
            if cv_experience_years is not None and job_experience_years is not None:
                experience_match = cv_experience_years >= job_experience_years
                experience_score = 100 if experience_match else 0

            # Địa điểm - lấy full location từ job nhưng so trùng theo city
            parts = []
            location_field = job.get("location")
            if isinstance(location_field, dict):
                parts.append(location_field.get("name", ""))
            elif isinstance(location_field, str):
                parts.append(location_field)
            parts.append(job.get("work_location_detail", ""))
            job_full_location = " | ".join([p for p in parts if p])
            job_city = extract_city(job_full_location) if job_full_location else ""
            location_match = check_location_match(cv_city, job_city) if cv_city and job_city else False
            location_score = 100 if location_match else (50 if (cv_city and job_city) else 0)

            # Tính match_percentage dựa trên 3 yếu tố: skill (70%), experience (15%), location (15%)
            match_percentage = round(skill_score * 0.7 + experience_score * 0.15 + location_score * 0.15, 1)

            # Tạo lý do giải thích
            reason_parts = []
            if len(matched_required) > 0:
                reason_parts.append(f"Match {len(matched_required)}/{len(job_required_skills)} kỹ năng bắt buộc")
            if len(matched_optional) > 0:
                reason_parts.append(f"Match {len(matched_optional)} kỹ năng tùy chọn")
            if experience_match is True:
                reason_parts.append("Kinh nghiệm đủ yêu cầu")
            elif experience_match is False:
                reason_parts.append("Kinh nghiệm có thể thấp hơn yêu cầu")
            if location_match:
                reason_parts.append("Địa điểm phù hợp")
            
            reason = "; ".join(reason_parts) if reason_parts else "Nội dung CV phù hợp với yêu cầu công việc"

            combined_score = float(score) + match_percentage / 100.0
            if experience_match is True:
                combined_score += 0.10
            if location_match:
                combined_score += 0.05

            results.append({
                "id": str(job["_id"]),
                "title": str(job.get("title", "")),
                "location": convert_objectid_to_str(job.get("location")),
                "salary": str(job.get("salary_raw", "")),
                "company": str(job.get("company", "")),
                "score": float(score),
                "similarity_percentage": round(float(score) * 100, 1),
                "match_percentage": match_percentage,
                "experience_required": job_experience_years,
                "experience_match": experience_match,
                "cv_experience_years": cv_experience_years,
                "location_match": location_match,
                "cv_location": cv_full_location,
                "cv_city": cv_city,
                "job_location_text": job_full_location,
                "job_city": job_city,
                "combined_score": round(combined_score, 4),
                "matched_skills": {
                    "required": list(matched_required),
                    "optional": list(matched_optional),
                    "domain": list(matched_domain),
                    "total_matched": len(matched_total),
                    "total_job_skills": len(job_all_skills)
                },
                "reason": reason
            })

        return convert_objectid_to_str({
            "skills_found": cv_skills,
            "recommendations": sorted(results, key=lambda x: x["combined_score"], reverse=True)[:100],
            "summary": f"Tìm được {len(cv_skills)} kỹ năng trong CV của bạn"
        })
    
    except Exception as e:
        print(f"Recommend error: {str(e)}")
        import traceback
        traceback.print_exc()
        return convert_objectid_to_str({
            "error": f"Lỗi xử lý: {str(e)}",
            "skills_found": [],
            "recommendations": []
        })
