from fastapi import FastAPI, UploadFile, File
import os
from fastapi.middleware.cors import CORSMiddleware
from io import BytesIO
import re
import pdfplumber
from docx import Document
import zipfile
from xml.etree import ElementTree as ET
import io
from sentence_transformers import SentenceTransformer
from rapidfuzz import fuzz
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.preprocessing import normalize
from pymongo import MongoClient
import numpy as np
from bson import ObjectId
from cities import get_canonical_city, is_city_match

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

model = SentenceTransformer("all-MiniLM-L6-v2")

client = MongoClient(os.getenv("MONGO_URL"))
db = client["ITJOBS"]
job_collection = db["jobs"]
skills_collection = db["skills"]

def convert_objectid_to_str(obj):
    """Recursively convert ObjectId to string for JSON serialization"""
    if isinstance(obj, ObjectId):
        return str(obj)
    elif isinstance(obj, dict):
        return {k: convert_objectid_to_str(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [convert_objectid_to_str(item) for item in obj]
    return obj


def load_skill_dict():
    skills = skills_collection.find({}, {"name": 1})
    return [s["name"].lower() for s in skills]

skill_dict = load_skill_dict()


def extract_pdf_text(file_like):
    text = ""
    with pdfplumber.open(file_like) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def extract_docx_text(file_like):
    """
    Extract text from DOCX by parsing Word XML parts (document, headers, footers,
    textboxes). This is more robust for CV templates that place contact info
    inside textboxes or headers/footers.
    """
    # Accept either raw bytes or a file-like object
    # Ensure we have a BytesIO
    if isinstance(file_like, (bytes, bytearray)):
        bio = io.BytesIO(file_like)
    else:
        try:
            # if file_like already a BytesIO or file object
            file_like.seek(0)
            bio = io.BytesIO(file_like.read())
        except Exception:
            # fallback: try to construct BytesIO directly
            bio = io.BytesIO()

    paragraphs = []

    try:
        with zipfile.ZipFile(bio) as z:
            # iterate relevant xml parts inside the docx package
            for name in z.namelist():
                if not name.startswith('word/') or not name.endswith('.xml'):
                    continue
                try:
                    xml = z.read(name)
                    root = ET.fromstring(xml)
                    # WordprocessingML namespace
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    for p in root.findall('.//w:p', ns):
                        texts = [t.text for t in p.findall('.//w:t', ns) if t.text]
                        if texts:
                            paragraphs.append(''.join(texts))
                except Exception:
                    # ignore parse errors for optional parts
                    continue
    except Exception:
        # fallback to python-docx if zip parsing fails
        try:
            doc = Document(io.BytesIO(bio.getvalue()))
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(' | '.join(row_text))
        except Exception:
            return ''

    return '\n'.join(paragraphs)

def extract_text_from_file(file_bytes, filename):
    filename = filename.lower()

    if filename.endswith(".pdf"):
        return extract_pdf_text(BytesIO(file_bytes))

    elif filename.endswith(".docx"):
        return extract_docx_text(BytesIO(file_bytes))

    else:
        raise ValueError("Chỉ hỗ trợ file PDF hoặc DOCX")


def extract_skills(text):
    text_lower = text.lower()
    skills = []
    for skill in skill_dict:
        if skill in text_lower or fuzz.partial_ratio(skill, text_lower) > 85:
            skills.append(skill)
    return list(set(skills))


def normalize_text(text):
    if not text:
        return ""
    return re.sub(r"\s+", " ", text.lower().strip())


def extract_city(text):
    """Extract city name from text using cities.py"""
    if not text:
        return ""
    return get_canonical_city(text)


def extract_location_text(text):
    """
    Extract full location/address from CV text.
    """
    if not text:
        return ""

    lines = [line.strip() for line in text.splitlines() if line.strip()]

    # helper checks
    address_keywords_regex = re.compile(r"\b(quận|q\.|phường|p\.|đường|duong|tp\.?|thành phố|thanh pho|tỉnh|district|city|street|st\.|road|rd\.|avenue|đ/c|đc)\b", flags=re.IGNORECASE)
    non_address_stopwords = re.compile(r"\b(managed|team|onshore|offshore|experience|working|years|year|project|projects|skills|responsible|have|and)\b", flags=re.IGNORECASE)

    # 1. Detect real address patterns first (priority): header-style US, VN, generic street+number
    address_patterns = [
        # US style with commas: 50 GRAHAM ST, JERSEY CITY, NJ 07307
        r"\d{1,5}\s+[A-Za-z0-9\.\- ]{3,},\s*[A-Za-z\s\.]{2,}(?:,|\s)\s*[A-Z]{2}\s+\d{4,6}",
        # US style without comma between city and state: 50 GRAHAM ST, JERSEY CITY NJ 07307
        r"\d{1,5}\s+[A-Za-z0-9\.\- ]{3,},\s*[A-Za-z\s\.]{2,}\s+[A-Z]{2}\s+\d{4,6}",
        # Uppercase header-like pattern: street + CITY STATE ZIP (no commas)
        r"\b\d{1,5}\s+[A-Z0-9\.\- ]{3,},?\s+[A-Z\s]{2,}\s+[A-Z]{2}\s*\d{5}\b",
        # Vietnam style: 123 Nguyễn Văn Cừ, Quận 5, TP.HCM
        r"\d+\s+[\w\sÀ-ỹ\.\-]+,\s*.*(?:quận|q\.|huyện|tp\.?|thành phố|district|city)",
        # Generic address line with number + street keywords
        r"\d+\s+.*(?:street|st\.?|road|rd\.?|avenue|ave\.?|boulevard|blvd\.?|lane|ln\.?)",
    ]
    # 2. Fallback: keyword-based markers (only whole-word regex matches, no short-substring fallback)
    markers = [
        "nơi ở",
        "noi o",
        "địa chỉ",
        "dia chi",
        "đ/c",
        "address",
        "location",
        "📍",
    ]

    for idx, line in enumerate(lines):
        lower = line.lower()

        for marker in markers:
            # match marker as separate word to avoid accidental matches inside other words
            try:
                m = re.search(rf"\b{re.escape(marker)}\b\s*[:\-–]?\s*(.*)", line, flags=re.IGNORECASE)
            except re.error:
                m = None

            if marker == "📍":
                if "📍" in line:
                    parts = line.split("📍", 1)
                    result = parts[1].strip() if len(parts) > 1 else ""
                else:
                    continue
            elif m:
                result = m.group(1).strip()
            else:
                continue

            # If marker was present but nothing meaningful after it, try next line as value
            if not result and idx + 1 < len(lines):
                next_line = lines[idx + 1].strip()
                result = next_line

            # clean leading separators
            result = re.sub(r"^[\-:\s]+", "", result)

            # Validation: accept only if it contains a recognized city, address keywords, or digits (street number)
            if extract_city(result) or address_keywords_regex.search(result) or re.search(r"\d{1,}", result):
                return result

            # Reject results that look like experience/role descriptions unless a city is found.
            if non_address_stopwords.search(result) and not extract_city(result):
                continue

            # If result looks reasonably like an address (multiple words, commas), accept it
            if len(result) >= 10 and ("," in result or len(result.split()) <= 6):
                return result

    # 2. Pattern-first pass with proximity to contact info and better substring extraction
    # address_patterns already defined above; compile hyphen pattern that captures only the hyphen-city fragment
    hyphen_re = re.compile(r"([\wÀ-ỹ0-9\s\.]{2,}?[-–—]\s*(?:tp\.?|thành phố|thanh pho|tp|tỉnh)[^,;\n\u2022\u2023\|]{0,80})", flags=re.IGNORECASE)

    # Contact proximity detection: prefer lines near email/phone/contact markers
    contact_re = re.compile(r"\S+@\S+|github\.com|linkedin\.com|\bthông tin\b|\bcontact\b|\bcontact:\b", flags=re.IGNORECASE)
    phone_re = re.compile(r"\+?\d[\d\-\.\s\(\)]{6,}\d")
    contact_indices = [i for i, L in enumerate(lines) if contact_re.search(L) or phone_re.search(L)]

    def city_snippet_from_line(line):
        low = line.lower()
        tokens = [
            'tp', 'thành phố', 'thanh pho', 'hồ chí minh', 'ho chi minh', 'hcm', 'sài gòn', 'saigon',
            'hanoi', 'thủ đức', 'thu duc', 'quận', 'quan', 'tỉnh', 'city', 'district'
        ]
        # prefer longer tokens (e.g., 'hồ chí minh') to avoid matching 'tp' first
        for tok in sorted(tokens, key=lambda x: -len(x)):
            pos = low.find(tok)
            if pos != -1:
                start = max(0, pos - 10)
                end = min(len(line), pos + len(tok) + 10)
                return line[start:end].strip(' ,;:-•\n\t—–')
        return line.strip()

    def trim_hyphen_match(s):
        low = s.lower()
        tokens_sorted = sorted([
            'tp', 'thành phố', 'thanh pho', 'hồ chí minh', 'ho chi minh', 'hcm', 'sài gòn', 'saigon',
            'hanoi', 'thủ đức', 'thu duc', 'quận', 'quan', 'tỉnh', 'city', 'district'
        ], key=lambda x: -len(x))

        for tok in tokens_sorted:
            pos = low.find(tok)
            if pos != -1:
                hyph = re.search(r"[-–—]", s)
                end = pos + len(tok)
                if hyph and hyph.start() < pos:
                    return s[:end].strip(' ,;:-•\n\t—–')
                return s[:end].strip(' ,;:-•\n\t—–')

        # fallback to city snippet
        return city_snippet_from_line(s)

    def is_education_or_org(line):
        low = line.lower()
        return bool(re.search(r"\b(đại học|dai hoc|sư phạm|sư phạm|university|college|institute|school|khoa|faculty)\b", low))

    def is_address_like(line):
        low = line.lower()
        if re.search(r"🏠|house|home|địa chỉ|dia chi|address|location", low):
            return True
        if re.search(r"[-–—]", line) and re.search(r"tp\.?|thành phố|thanh pho|quận|quan|p\.|phường", low):
            return True
        if re.search(r"\d+\s+", line):
            return True
        return False

    # Check near contact info first
    for idx in contact_indices:
        for j in range(idx - 2, idx + 4):
            if j < 0 or j >= len(lines):
                continue
            cand = lines[j]
            cleaned = re.sub(r"\S+@\S+", "", cand)
            cleaned = re.sub(r"[•\u2022\u2023\|\u25CF]", " ", cleaned)
            cleaned = re.sub(r"(\+?\d[\d\-\.\s\(\)]{6,}\d)", " ", cleaned)
            cleaned = cleaned.strip()
            if not cleaned:
                continue

            m = hyphen_re.search(cleaned)
            if m:
                return trim_hyphen_match(m.group(1).strip())

            for pat in address_patterns:
                mm = re.search(pat, cleaned, re.IGNORECASE)
                if mm:
                    return mm.group(0).strip()

                if extract_city(cleaned):
                    # accept city-only line near contact only if it looks like an address and not an education/org line
                    if is_address_like(cleaned) and not is_education_or_org(cleaned):
                        return city_snippet_from_line(cleaned)
                    # also accept very short lines (<=6 words) that are not organization names
                    words = cleaned.split()
                    if len(words) <= 6 and not is_education_or_org(cleaned):
                        return city_snippet_from_line(cleaned)

    # If nothing found near contacts, do line-by-line pattern-first extraction
    for line in lines:
        cleaned = re.sub(r"\S+@\S+", "", line)
        cleaned = re.sub(r"[•\u2022\u2023\|\u25CF]", " ", cleaned)
        cleaned = re.sub(r"(\+?\d[\d\-\.\s\(\)]{6,}\d)", " ", cleaned)
        cleaned = cleaned.strip()
        if not cleaned:
            continue

        m = hyphen_re.search(cleaned)
        if m:
            return trim_hyphen_match(m.group(1).strip())

        for pat in address_patterns:
            mm = re.search(pat, cleaned, re.IGNORECASE)
            if mm:
                return mm.group(0).strip()
    # Try combining top header lines (some CVs split address across inline header)
    header_candidate = " ".join(lines[:3])
    header_cleaned = re.sub(r"\S+@\S+", "", header_candidate)
    header_cleaned = re.sub(r"(\+?\d[\d\-\.\s\(\)]{6,}\d)", " ", header_cleaned)
    header_cleaned = re.sub(r"[•\u2022\u2023\|\u25CF]", " ", header_cleaned).strip()
    m = hyphen_re.search(header_cleaned)
    if m:
        return trim_hyphen_match(m.group(1).strip())
    for pat in address_patterns:
        mm = re.search(pat, header_cleaned, re.IGNORECASE)
        if mm:
            return mm.group(0).strip()

    # Fallback by city detection: return a short snippet around the city token
    for line in lines:
        if extract_city(line):
            # final fallback: only return if the line is address-like and not an education/org
            if is_address_like(line) and not is_education_or_org(line):
                return city_snippet_from_line(line)

    return ""


def parse_experience_text(text):
    if not text:
        return None

    normalized = text.lower()
    no_experience_markers = [
        "no formal work experience",
        "no experience",
        "chưa có kinh nghiệm",
        "không có kinh nghiệm",
        "no working experience",
        "no professional experience",
    ]
    for marker in no_experience_markers:
        if marker in normalized:
            return 0

    def parse_experience_line(line):
        if not line:
            return None
        line_lower = line.lower()
        range_match = re.search(
            r"(\d+)\s*(?:\+|\-\s*|to|đến|den|–)\s*(\d+)\s*(?:năm|nam|year|years|yr|yrs)",
            line_lower,
        )
        if range_match:
            return max(int(range_match.group(1)), int(range_match.group(2)))
        single_match = re.search(r"(\d+)\s*(?:\+)?\s*(?:năm|nam|year|years|yr|yrs)", line_lower)
        if single_match:
            return int(single_match.group(1))
        return None

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    experience_lines = []
    for line in lines:
        lower_line = line.lower()
        if any(marker in lower_line for marker in [
            "kinh nghiệm",
            "experience",
            "work experience",
            "years of experience",
            "years experience",
            "năm kinh nghiệm",
        ]):
            experience_lines.append(line)

    if not experience_lines:
        for index, line in enumerate(lines):
            lower_line = line.lower()
            if any(marker in lower_line for marker in ["kinh nghiệm", "experience"]):
                if index + 1 < len(lines):
                    experience_lines.append(lines[index + 1])

    for line in experience_lines:
        for marker in no_experience_markers:
            if marker in line.lower():
                return 0
        parsed = parse_experience_line(line)
        if parsed is not None:
            return parsed

    return None


def parse_job_experience(experience):
    if experience is None:
        return None
    if isinstance(experience, (int, float)):
        return int(experience)
    text = str(experience).lower()
    if not text:
        return None
    range_match = re.search(r"(\d+)\s*(?:\+|\-\s*|to|đến|den|–)\s*(\d+)\s*(?:năm|nam|year|years|yr|yrs)?", text)
    if range_match:
        return max(int(range_match.group(1)), int(range_match.group(2)))
    single_match = re.search(r"(\d+)\s*(?:\+)?\s*(?:năm|nam|year|years|yr|yrs)", text)
    if single_match:
        return int(single_match.group(1))
    digits = re.findall(r"\d+", text)
    return int(digits[0]) if digits else None


def get_job_location_city(job):
    """Extract city from job location fields"""
    parts = []
    location_field = job.get("location")
    if isinstance(location_field, dict):
        parts.append(location_field.get("name", ""))
    elif isinstance(location_field, str):
        parts.append(location_field)
    parts.append(job.get("work_location_detail", ""))
    combined = " ".join([p for p in parts if p])
    return extract_city(combined)


def check_location_match(cv_location, job_location):
    """Check if CV location and job location are in same city"""
    if not cv_location or not job_location:
        return False
    # Use cities.py logic for city matching
    return is_city_match(cv_location, job_location)


def generate_job_embedding(text):
    """Generate embedding from text when job doesn't have pre-computed embedding"""
    try:
        emb_full = model.encode(text)
        return emb_full.tolist()
    except:
        return np.zeros(384).tolist()


def prepare_job_skills(job):
    """Extract job skills with fallback for old jobs without structured fields"""
    # Prioritize structured fields
    must_have = job.get("mustHaveSkills", [])
    optional = job.get("optionalSkills", [])
    domain = job.get("domainKnowledge", [])
    
    # Fallback: extract from requirements text if structured fields are empty
    if not must_have and not optional and not domain:
        requirements_text = " ".join(job.get("requirements", []))
        if not requirements_text:
            requirements_text = job.get("description", "")
        
        if requirements_text:
            extracted = extract_skills(requirements_text)
            # Treat all extracted skills as must-have for old jobs
            must_have = extracted
    
    return {
        "must_have": set([s.lower() for s in must_have]) if must_have else set(),
        "optional": set([s.lower() for s in optional]) if optional else set(),
        "domain": set([s.lower() for s in domain]) if domain else set()
    }


def generate_cv_embedding(raw_text, skills):
    skills_text = " ".join(skills) if skills else ""

    emb_skill = (
        model.encode(skills_text)
        if skills_text
        else np.zeros(384)
    )

    emb_full = model.encode(raw_text)

    return (
        0.7 * emb_skill +
        0.3 * emb_full
    ).tolist()
    